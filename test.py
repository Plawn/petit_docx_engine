import docx
from docx.document import Document


doc: Document = docx.Document('cet_hec.docx')


for section in doc.sections:
    print(section)
    print(type(section), dir(section))
    header = section.header
    footer = section.footer
    for paragraph in header.paragraphs:
        print(paragraph.text)  # or whatever you have in mind

    print(footer.paragraphs)
    obj = footer._document_part
    print(type(obj), dir(obj))
    for paragraph in footer.paragraphs:
        print(paragraph)
        print(paragraph.text)  # or whatever you have in mind

    print('first')
    footer = section.first_page_footer
    for paragraph in footer.paragraphs:
        print(paragraph.text)  # or whatever you have in mind
    
    print('first')
    header = section.first_page_header
    for paragraph in header.paragraphs:
        print(paragraph.text)  # or whatever you have in mind
        # print(paragraph.part)  # or whatever you have in mind

    print('even')
    footer = section.even_page_footer
    for paragraph in footer.paragraphs:
        print('text=', paragraph.text)  # or whatever you have in mind

    print('even')
    header = section.even_page_header
    for paragraph in header.paragraphs:
        print('text=', paragraph.text)  # or whatever you have in mind

    for table in footer.tables:
        for row in table.rows:
            for cell in row.cells:
                print('text=', cell.text)  # or whatever you have in mind

    print('footer')
    footer = section.footer
    for paragraph in footer.paragraphs:
        print(paragraph.text)  # or whatever you have in mind


# print(c_doc.fields)
